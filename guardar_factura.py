from flask import Flask, request, jsonify, render_template
from docx import Document
from docx2pdf import convert
import os
import tempfile
from datetime import datetime
import pythoncom

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/guardar_pdf', methods=['POST'])
def guardar_pdf():
    try:
        data = request.json

        plantilla_path = 'plantilla_factura.docx'
        doc = Document(plantilla_path)

        # Reemplazo de campos excepto artículos
        for p in doc.paragraphs:
            for key, value in data.items():
                if key == "articulos":
                    continue
                placeholder = f'{{{{{key}}}}}'
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, str(value))

        # Reemplazo de tabla_articulos
        for i, p in enumerate(doc.paragraphs):
            if "{{tabla_articulos}}" in p.text:
                p.text = ""
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Artículo'
                hdr_cells[1].text = 'Cantidad'
                hdr_cells[2].text = 'Precio'

                for art in data.get("articulos", []):
                    row_cells = table.add_row().cells
                    row_cells[0].text = art.get('articulo', '')
                    row_cells[1].text = art.get('cantidad', '')
                    row_cells[2].text = art.get('precio', '')
                
                p._element.addnext(table._element)
                break

        # Gestión automática de archivos temporales
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_docx = os.path.join(temp_dir, 'factura.docx')
            doc.save(temp_docx)

            escritorio = os.path.join(os.path.expanduser("~"), "Desktop")
            pdf_name = f"factura_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            salida_pdf = os.path.join(escritorio, pdf_name)

            # Convertir a PDF
            pythoncom.CoInitialize()  # <-- Inicializa COM
            convert(temp_docx, salida_pdf)

        return jsonify({'message': f'Factura guardada en {salida_pdf}'})

    except Exception as e:
        return jsonify({'message': f'Error al guardar la factura: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True)
